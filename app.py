import re
import json
from datetime import datetime
import streamlit as st
from openai import OpenAI
from openai import RateLimitError, APIError
from docx import Document

st.set_page_config(page_title="Listing Marketing Assistant", page_icon="🏡", layout="wide")

# -----------------------
# Simple password gate (optional)
# -----------------------
def require_password() -> None:
    expected = st.secrets.get("APP_PASSWORD", "")
    if not expected:
        return

    if "authed" not in st.session_state:
        st.session_state.authed = False

    if st.session_state.authed:
        return

    st.title("🏡 Listing Marketing Assistant")
    st.caption("Enter the demo password to continue.")
    pwd = st.text_input("Password", type="password")

    if st.button("Unlock"):
        if pwd == expected:
            st.session_state.authed = True
            st.rerun()
        else:
            st.error("Incorrect password.")
    st.stop()

require_password()

# -----------------------
# Client setup
# -----------------------
api_key = st.secrets.get("OPENAI_API_KEY")
if not api_key:
    st.error("Missing OPENAI_API_KEY in Streamlit Secrets.")
    st.stop()

client = OpenAI(api_key=api_key)

# -----------------------
# Session state
# -----------------------
if "history" not in st.session_state:
    st.session_state.history = []
if "brand_voice" not in st.session_state:
    st.session_state.brand_voice = (
        "Friendly, clear, professional. Avoid hype and ALL CAPS. "
        "Focus on benefits, key features, and accuracy."
    )
if "busy" not in st.session_state:
    st.session_state.busy = False


# -----------------------
# Helpers
# -----------------------
def clean_text(s: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", (s or "").strip())

def safe_parse_json(text: str) -> dict:
    try:
        return json.loads(text)
    except Exception:
        start, end = text.find("{"), text.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                return json.loads(text[start:end+1])
            except Exception:
                pass
    return {}

def make_docx(title: str, meta: dict, outputs: dict) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    if meta:
        doc.add_paragraph(f"Address: {meta.get('address','')}")
        if meta.get("price"):
            doc.add_paragraph(f"Price: {meta.get('price')}")
        stats = f"{meta.get('beds','')} bd • {meta.get('baths','')} ba • {meta.get('sqft','')} sqft"
        doc.add_paragraph(stats)
        if meta.get("listing_url"):
            doc.add_paragraph(f"Listing URL: {meta.get('listing_url')}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_paragraph("")

    mls = outputs.get("mls", {})
    social = outputs.get("social", {})
    email = outputs.get("email", {})

    doc.add_heading("MLS Listing", level=2)
    doc.add_paragraph(mls.get("headline", ""))
    doc.add_paragraph(mls.get("description", ""))

    doc.add_heading("Social Media", level=2)
    doc.add_heading("Instagram", level=3)
    doc.add_paragraph(social.get("instagram", ""))
    doc.add_heading("Facebook", level=3)
    doc.add_paragraph(social.get("facebook", ""))
    tags = social.get("hashtags", [])
    if tags:
        doc.add_paragraph("Hashtags: " + " ".join([t if t.startswith("#") else f"#{t}" for t in tags]))

    doc.add_heading("Buyer Email", level=2)
    doc.add_paragraph("Subject: " + email.get("subject", ""))
    doc.add_paragraph(email.get("body", ""))

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_generation_prompt(meta: dict) -> str:
    voice = st.session_state.brand_voice
    tone = meta.get("tone_preset", "Professional MLS")
    include_cta = meta.get("include_cta", True)
    cta_rule = "Include a clear call-to-action." if include_cta else "No call-to-action."

    return f"""
You are a real estate marketing expert and MLS-compliant copywriter.

Brokerage voice profile:
{voice}

Rules:
- Avoid discriminatory language or references to protected classes.
- Keep MLS description factual, not exaggerated.
- Do not invent details not provided.
- If something is missing (e.g., price), omit it gracefully.

Task:
Create:
1) MLS headline (max 12 words) + MLS description (120–180 words)
2) Social pack: Instagram + Facebook + 5 hashtags
3) Buyer email: subject + 5–8 sentences

Tone preset: {tone}
CTA rule: {cta_rule}

Property details:
Address: {meta.get("address","")}
Price: {meta.get("price","")}
Beds: {meta.get("beds","")}
Baths: {meta.get("baths","")}
Square Feet: {meta.get("sqft","")}
Highlights: {meta.get("highlights","")}
Neighborhood notes: {meta.get("neighborhood","")}

Return STRICT JSON with schema:
{{
  "mls": {{
    "headline": "string",
    "description": "string"
  }},
  "social": {{
    "instagram": "string",
    "facebook": "string",
    "hashtags": ["string","string","string","string","string"]
  }},
  "email": {{
    "subject": "string",
    "body": "string"
  }}
}}
"""

def build_extract_prompt(listing_text: str) -> str:
    return f"""
Extract real-estate listing details from the text below.

Return STRICT JSON with schema:
{{
  "address": "string or empty",
  "price": "string or empty",
  "beds": "number or empty",
  "baths": "number or empty",
  "sqft": "number or empty",
  "highlights": "string (comma-separated features)",
  "neighborhood": "string"
}}

Text:
{listing_text}
"""

def call_model(prompt: str, max_tokens: int = 800, temperature: float = 0.7) -> str:
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "Return exactly what the user asked for. If JSON requested, output valid JSON only."},
            {"role": "user", "content": prompt},
        ],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return resp.choices[0].message.content


# -----------------------
# UI (friendly)
# -----------------------
st.title("🏡 Listing Marketing Assistant")
st.caption("Paste a listing link (optional) + details. Generate MLS copy, social posts, and a buyer email.")

left, right = st.columns([0.95, 1.05], gap="large")

with left:
    st.subheader("1) Add listing info")

    with st.expander("🏢 Brokerage voice (saved)", expanded=False):
        st.session_state.brand_voice = st.text_area(
            "How should it sound?",
            value=st.session_state.brand_voice,
            height=110
        )

    listing_url = st.text_input("Listing URL (Zillow/Redfin/any)", placeholder="https://www.zillow.com/homedetails/...")
    st.caption("Tip: For Zillow, paste the link for reference, then copy/paste the listing text below for best results.")

    pasted_listing_text = st.text_area(
        "Paste listing text (Facts & Features / Description)",
        placeholder="Copy from Zillow/MLS: beds, baths, sqft, upgrades, neighborhood notes…",
        height=160
    )

    colx, coly = st.columns(2)
    with colx:
        tone_preset = st.selectbox(
            "Tone",
            ["Professional MLS", "Warm + inviting", "Modern + punchy", "Luxury", "Investor-focused"],
            index=0
        )
    with coly:
        include_cta = st.checkbox("Include a call-to-action", value=True)

    st.markdown("---")
    st.subheader("2) Or type manually (optional)")

    with st.form("manual_form"):
        address = st.text_input("Address", placeholder="123 Maple St, Toledo, OH")
        price = st.text_input("Price", placeholder="$325,000")

        c1, c2, c3 = st.columns(3)
        with c1:
            beds = st.number_input("Bedrooms", min_value=0, max_value=20, value=3)
        with c2:
            baths = st.number_input("Bathrooms", min_value=0.0, max_value=20.0, value=2.0, step=0.5)
        with c3:
            sqft = st.number_input("Square feet", min_value=0, max_value=20000, value=1850, step=50)

        highlights = st.text_area("Highlights / Features", height=90)
        neighborhood = st.text_area("Neighborhood / Location notes", height=70)

        submitted_manual = st.form_submit_button("Generate marketing pack", disabled=st.session_state.busy)

    col_extract, col_generate = st.columns(2)
    with col_extract:
        extract_clicked = st.button("✨ Fill fields from pasted text", disabled=st.session_state.busy)
    with col_generate:
        generate_from_paste_clicked = st.button("Generate from pasted text", disabled=st.session_state.busy)

    # Store the latest "working meta" dict here
    if "working_meta" not in st.session_state:
        st.session_state.working_meta = {}

    # Extraction step
    if extract_clicked:
        if not pasted_listing_text.strip():
            st.warning("Paste listing text first.")
        else:
            st.session_state.busy = True
            try:
                with st.spinner("Extracting details..."):
                    raw = call_model(build_extract_prompt(pasted_listing_text), max_tokens=350, temperature=0.2)
                    data = safe_parse_json(raw) or {}
                # Merge into manual fields (best-effort)
                st.session_state.working_meta = {
                    "address": data.get("address", ""),
                    "price": data.get("price", ""),
                    "beds": data.get("beds", ""),
                    "baths": data.get("baths", ""),
                    "sqft": data.get("sqft", ""),
                    "highlights": data.get("highlights", ""),
                    "neighborhood": data.get("neighborhood", ""),
                    "listing_url": listing_url,
                    "tone_preset": tone_preset,
                    "include_cta": include_cta,
                }
                st.success("Extracted! Scroll down on the right to generate, or re-run with tweaks.")
            except RateLimitError:
                st.error("Rate limit / quota hit. Try again in a moment.")
            except APIError as e:
                st.error(f"OpenAI API error: {e}")
            except Exception as e:
                st.error(f"Unexpected error: {e}")
            finally:
                st.session_state.busy = False

    # Generate directly from pasted text
    if generate_from_paste_clicked:
        if not pasted_listing_text.strip():
            st.warning("Paste listing text first.")
        else:
            st.session_state.busy = True
            try:
                with st.spinner("Generating marketing pack..."):
                    # Extract first (quick) then generate
                    raw = call_model(build_extract_prompt(pasted_listing_text), max_tokens=350, temperature=0.2)
                    data = safe_parse_json(raw) or {}
                    meta = {
                        "address": data.get("address", ""),
                        "price": data.get("price", ""),
                        "beds": data.get("beds", ""),
                        "baths": data.get("baths", ""),
                        "sqft": data.get("sqft", ""),
                        "highlights": data.get("highlights", ""),
                        "neighborhood": data.get("neighborhood", ""),
                        "listing_url": listing_url,
                        "tone_preset": tone_preset,
                        "include_cta": include_cta,
                    }
                    out_raw = call_model(build_generation_prompt(meta), max_tokens=900, temperature=0.7)
                    outputs = safe_parse_json(out_raw) or {"_raw": out_raw}

                st.session_state.history.insert(0, {"meta": meta, "outputs": outputs})
                st.success("Done! See results on the right →")
            except RateLimitError:
                st.error("Rate limit / quota hit. Try again in a moment.")
            except APIError as e:
                st.error(f"OpenAI API error: {e}")
            except Exception as e:
                st.error(f"Unexpected error: {e}")
            finally:
                st.session_state.busy = False

    # Generate from manual form
    if submitted_manual:
        meta = {
            "address": address,
            "price": price,
            "beds": beds,
            "baths": baths,
            "sqft": sqft,
            "highlights": highlights,
            "neighborhood": neighborhood,
            "listing_url": listing_url,
            "tone_preset": tone_preset,
            "include_cta": include_cta,
        }
        if not meta["address"] and not meta["highlights"]:
            st.warning("Add an address or highlights so the copy has something to work with.")
        else:
            st.session_state.busy = True
            try:
                with st.spinner("Generating marketing pack..."):
                    out_raw = call_model(build_generation_prompt(meta), max_tokens=900, temperature=0.7)
                    outputs = safe_parse_json(out_raw) or {"_raw": out_raw}
                st.session_state.history.insert(0, {"meta": meta, "outputs": outputs})
                st.success("Done! See results on the right →")
            except RateLimitError:
                st.error("Rate limit / quota hit. Try again in a moment.")
            except APIError as e:
                st.error(f"OpenAI API error: {e}")
            except Exception as e:
                st.error(f"Unexpected error: {e}")
            finally:
                st.session_state.busy = False


with right:
    st.subheader("Marketing Pack")

    if not st.session_state.history:
        st.info("Generate a marketing pack and it will show up here.")
    else:
        # Picker
        labels = []
        for i, item in enumerate(st.session_state.history):
            m = item["meta"]
            addr = m.get("address") or "(no address)"
            labels.append(f"{i+1}. {addr} — {m.get('price','')}")

        selected = st.selectbox("Saved packs", labels, index=0)
        idx = int(selected.split(".")[0]) - 1
        item = st.session_state.history[idx]

        meta = item["meta"]
        outputs = item["outputs"]

        st.markdown(
            f"**{meta.get('address','')}**  \n"
            f"{meta.get('beds','')} bd • {meta.get('baths','')} ba • {meta.get('sqft','')} sqft"
            + (f" • **{meta.get('price','')}**" if meta.get("price") else "")
        )
        if meta.get("listing_url"):
            st.caption(f"Link: {meta.get('listing_url')}")

        mls = outputs.get("mls", {})
        social = outputs.get("social", {})
        email = outputs.get("email", {})

        tab1, tab2, tab3 = st.tabs(["MLS Listing", "Social Posts", "Buyer Email"])

        with tab1:
            st.text_input("MLS Headline", value=clean_text(mls.get("headline","")), key=f"headline_{idx}")
            st.text_area(
                "MLS Description",
                value=clean_text(mls.get("description","") or outputs.get("_raw","")),
                height=220,
                key=f"mlsdesc_{idx}"
            )

        with tab2:
            st.text_area("Instagram", value=clean_text(social.get("instagram","")), height=140, key=f"ig_{idx}")
            st.text_area("Facebook", value=clean_text(social.get("facebook","")), height=140, key=f"fb_{idx}")
            tags = social.get("hashtags", [])
            if tags:
                tag_line = " ".join([t if t.startswith("#") else f"#{t.replace(' ','')}" for t in tags])
                st.text_input("Hashtags", value=tag_line, key=f"tags_{idx}")

        with tab3:
            st.text_input("Email Subject", value=clean_text(email.get("subject","")), key=f"subj_{idx}")
            st.text_area("Email Body", value=clean_text(email.get("body","")), height=220, key=f"body_{idx}")

        # Friendly downloads
        combined_txt = "\n\n".join([
            "MLS HEADLINE:\n" + clean_text(mls.get("headline","")),
            "MLS DESCRIPTION:\n" + clean_text(mls.get("description","") or outputs.get("_raw","")),
            "INSTAGRAM:\n" + clean_text(social.get("instagram","")),
            "FACEBOOK:\n" + clean_text(social.get("facebook","")),
            "HASHTAGS:\n" + (" ".join(social.get("hashtags", [])) if social.get("hashtags") else ""),
            "EMAIL SUBJECT:\n" + clean_text(email.get("subject","")),
            "EMAIL BODY:\n" + clean_text(email.get("body","")),
        ]).strip()

        col_d1, col_d2, col_d3 = st.columns([1, 1, 1])
        with col_d1:
            st.download_button(
                "Download Marketing Pack (.txt)",
                data=combined_txt.encode("utf-8"),
                file_name="marketing-pack.txt",
                mime="text/plain"
            )
        with col_d2:
            docx_bytes = make_docx("Marketing Pack", meta, outputs)
            st.download_button(
                "Download as Word (.docx)",
                data=docx_bytes,
                file_name="marketing-pack.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col_d3:
            if st.button("Delete this pack"):
                st.session_state.history.pop(idx)
                st.rerun()

st.caption("Note: Zillow links are accepted for reference; for best results paste the listing text into the app.")
